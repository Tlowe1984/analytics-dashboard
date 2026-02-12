#!/usr/bin/env python3
"""
Parse WXX Health Canonical Program Review document
Extracts Wins, Exec Summary sections (text bullets) and Decisions section (table)
"""

import sys
import json
from docx import Document
from docx.table import Table
from rich_text_parser_v2 import extract_rich_text

def parse_hearing_review(docx_path):
    """
    Parse Health Canonical Program Review document to extract:
    - Wins: bullet points under 🏆Wins or Wins heading
    - Exec Summary: bullet points under 🚀Exec Summary or Exec Summary heading
    - Decisions: table rows from WW Health Decisions table
    Returns structured data with rich text formatting preserved.
    """
    doc = Document(docx_path)
    
    items = []
    current_section = None
    order = 0
    in_decisions_table = False
    seen_fitness_algos = False  # Flag to stop after Fitness Algos
    
    # Process all document elements (paragraphs and tables)
    for element in doc.element.body:
        # Handle paragraphs
        if element.tag.endswith('p'):
            # Find the paragraph object
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            
            if para is None:
                continue
                
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if this is a section header (with or without emoji)
            text_lower = text.lower()
            
            # Look for Wins section (🏆Wins or just Wins)
            if 'wins' in text_lower and len(text) < 20:
                current_section = 'wins'
                in_decisions_table = False
                continue
            
            # Look for Exec Summary section (🚀Exec Summary or just Exec Summary)
            if 'exec summary' in text_lower and len(text) < 30:
                current_section = 'exec_summary'
                in_decisions_table = False
                continue
            
            # Stop Exec Summary section at status indicator lines or subsection headers
            if current_section == 'exec_summary':
                # Check for status indicator pattern [RED] [ORANGE] [YELLOW] [GREEN]
                if '[RED]' in text or '[ORANGE]' in text or '[YELLOW]' in text or '[GREEN]' in text:
                    if text.count('[') >= 3:  # Multiple status indicators
                        current_section = None
                        continue
                
                # Check for subsection headers that end Exec Summary
                if any(keyword in text_lower for keyword in ['experience overall', 'risks/opens', 'help needed']):
                    if len(text) < 50:  # Likely a header
                        current_section = None
                        continue
            
            # Look for Help Needed section header - skip this section entirely
            if 'help needed' in text_lower or 'flags for leadership' in text_lower:
                current_section = 'help_needed'
                continue
            
            # Look for Decisions section header
            if 'decisions' in text_lower and len(text) < 30:
                current_section = 'decisions'
                in_decisions_table = False
                continue
            
            # Skip if we're in decisions or help_needed section waiting for table
            if current_section in ['decisions', 'help_needed']:
                continue
            
            # Skip if we haven't found a section yet
            if current_section is None:
                continue
            
            # Stop exec_summary section immediately after Fitness Algos
            if current_section == 'exec_summary' and seen_fitness_algos:
                current_section = None
                continue
            
            # Extract rich text content for Wins and Exec Summary
            rich_text = extract_rich_text(para)
            
            if not rich_text or rich_text.strip() == '':
                continue
            
            # Check if this is the Fitness Algos bullet
            if current_section == 'exec_summary' and 'Fitness Algos' in rich_text:
                seen_fitness_algos = True
            
            # Detect if this is a "new" item (blue text)
            is_new = False
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    r, g, b = run.font.color.rgb
                    # Blue text detection
                    if b > 150 and r < 100 and g < 100:
                        is_new = True
                        break
            
            # Detect and strip [wearables-tag]
            is_wearables_tag = '[wearables-tag]' in rich_text.lower()
            if is_wearables_tag:
                # Remove [wearables-tag] (case insensitive)
                import re
                rich_text = re.sub(r'\[wearables-tag\]', '', rich_text, flags=re.IGNORECASE).strip()
            
            # Detect indent level
            indent_level = 0
            if para.paragraph_format.left_indent:
                indent_inches = para.paragraph_format.left_indent.inches
                indent_level = int(indent_inches / 0.5)
            
            items.append({
                'section_type': current_section,
                'content': rich_text,
                'is_new': is_new,
                'is_wearables_tag': is_wearables_tag,
                'indent_level': indent_level,
                'order': order
            })
            
            order += 1
        
        # Handle tables (for Decisions section)
        elif element.tag.endswith('tbl'):
            # Find the table object
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break
            
            if table is None:
                continue
            
            # Only process tables with "Health Decisions" in the header (regardless of section)
            if len(table.rows) == 0:
                continue
            
            first_cell_text = table.rows[0].cells[0].text.strip()
            if 'Health Decisions' not in first_cell_text:
                continue  # Skip tables that aren't the main Health Decisions table
            
            # Parse table rows (skip title row and header row)
            for i, row in enumerate(table.rows):
                if i == 0 or i == 1:  # Skip title row (0) and header row (1)
                    continue
                
                cells = row.cells
                if len(cells) < 6:  # Need at least 6 columns
                    continue
                
                # Extract cell contents with rich text
                topic = extract_rich_text(cells[0].paragraphs[0]) if cells[0].paragraphs else cells[0].text.strip()
                date = cells[1].text.strip()
                dri = cells[2].text.strip()
                forum = cells[3].text.strip()
                status = cells[4].text.strip()
                decision_doc = extract_rich_text(cells[5].paragraphs[0]) if cells[5].paragraphs else cells[5].text.strip()
                decision_makers = cells[6].text.strip() if len(cells) > 6 else ''
                
                # Skip empty rows
                if not topic or topic == '':
                    continue
                
                # Detect and strip [wearables-tag] from topic
                is_wearables_tag = '[wearables-tag]' in topic.lower()
                if is_wearables_tag:
                    import re
                    topic = re.sub(r'\[wearables-tag\]', '', topic, flags=re.IGNORECASE).strip()
                
                # Create decision item with table structure
                items.append({
                    'section_type': 'decisions',
                    'content': topic,  # Main content is the topic
                    'date': date,
                    'dri': dri,
                    'forum': forum,
                    'status': status,
                    'decision_doc': decision_doc,
                    'decision_makers': decision_makers,
                    'is_new': False,
                    'is_wearables_tag': is_wearables_tag,
                    'indent_level': 0,
                    'order': order
                })
                
                order += 1
    
    return {
        'items': items,
        'total_count': len(items)
    }

def main():
    if len(sys.argv) < 2:
        print(json.dumps({'error': 'Usage: parse_hearing_review.py <docx_path>'}))
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    try:
        result = parse_hearing_review(docx_path)
        print(json.dumps(result, indent=2))
    except Exception as e:
        print(json.dumps({'error': str(e)}), file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main()
