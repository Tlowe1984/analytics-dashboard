#!/usr/bin/python3.11
"""
Parse Software (I+E, AI, Hearing) Canonical Program Review document
Extracts Wins, Exec Summary, and Product Decisions table
"""

import sys
import json
from docx import Document
from docx.shared import RGBColor
from rich_text_parser_v2 import extract_rich_text

def is_blue_text(run):
    """Check if text run has blue color (indicating new information)"""
    if run.font.color and run.font.color.rgb:
        r, g, b = run.font.color.rgb
        # Check if color is predominantly blue
        return b > 150 and b > r and b > g
    return False

def extract_table_text(cell):
    """Extract text from a table cell, preserving rich text"""
    if not cell.paragraphs:
        return ""
    
    # Combine all paragraphs in the cell
    parts = []
    for para in cell.paragraphs:
        rich_text = extract_rich_text(para)
        if rich_text.strip():
            parts.append(rich_text)
    
    return " ".join(parts).strip()

def parse_product_decisions_table(doc):
    """
    Extract Product Decisions table from the document
    Returns list of decision items
    """
    decisions = []
    found_product_decisions = False
    current_category = None
    
    # First, find where "Product Decisions" appears in the document
    for i, para in enumerate(doc.paragraphs):
        if "Product Decisions" in para.text:
            found_product_decisions = True
            break
    
    if not found_product_decisions:
        return decisions
    
    # Now look at tables and find ones that have decision-related headers
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        
        # Check if first row is a title row (all cells have same text)
        first_row_cells = [cell.text.strip() for cell in table.rows[0].cells]
        is_title_row = len(set(first_row_cells)) == 1 and first_row_cells[0]
        
        # Get header row (skip title row if present)
        header_row_idx = 1 if is_title_row else 0
        if len(table.rows) <= header_row_idx:
            continue
        
        header_row = table.rows[header_row_idx]
        headers = [cell.text.strip().lower() for cell in header_row.cells]
        
        # Check if this looks like a decisions table
        has_topic = any("topic" in h for h in headers)
        has_status = any("status" in h for h in headers)
        has_outcome = any("outcome" in h for h in headers)
        
        if not (has_topic or has_status or has_outcome):
            continue  # Not a decisions table
        
        # Find column indices
        col_indices = {}
        for i, header in enumerate(headers):
            if "topic" in header:
                col_indices['topic'] = i
            elif "dri" in header:
                col_indices['dri'] = i
            elif "forum" in header:
                col_indices['forum'] = i
            elif "status" in header:
                col_indices['status'] = i
            elif "decision doc" in header or ("decision" in header and "doc" in header):
                col_indices['decision_doc'] = i
            elif "decision makers" in header or "reviewers" in header or "makers" in header:
                col_indices['decision_makers'] = i
            elif "decision outcome" in header or "outcome" in header:
                col_indices['decision_outcome'] = i
            elif "post" in header:
                col_indices['post'] = i
        
        # Parse data rows (skip title + header rows)
        data_start_idx = header_row_idx + 1
        for row_idx, row in enumerate(table.rows[data_start_idx:]):
            if len(row.cells) < 3:  # Skip invalid rows
                continue
            
            # Extract cell values safely
            topic = extract_table_text(row.cells[col_indices.get('topic', 0)]) if 'topic' in col_indices else ""
            
            # Check if this is a category header row
            topic_lower = topic.lower().replace("*", "").strip()
            if "pillar decisions" in topic_lower or "fyi sub-pillar" in topic_lower:
                # Update category and skip this row
                if "fyi" in topic_lower:
                    current_category = "FYI"
                else:
                    current_category = "Pillar"
                continue
            
            # Skip empty rows
            if not topic:
                continue
            
            dri = extract_table_text(row.cells[col_indices.get('dri', 1)]) if 'dri' in col_indices else ""
            forum = extract_table_text(row.cells[col_indices.get('forum', 2)]) if 'forum' in col_indices else ""
            status = extract_table_text(row.cells[col_indices.get('status', 3)]) if 'status' in col_indices else ""
            decision_doc = extract_table_text(row.cells[col_indices.get('decision_doc', 4)]) if 'decision_doc' in col_indices else ""
            decision_makers = extract_table_text(row.cells[col_indices.get('decision_makers', 5)]) if 'decision_makers' in col_indices else ""
            decision_outcome = extract_table_text(row.cells[col_indices.get('decision_outcome', 6)]) if 'decision_outcome' in col_indices else ""
            post = extract_table_text(row.cells[col_indices.get('post', 7)]) if 'post' in col_indices else ""
            
            decisions.append({
                "section_type": "decisions",
                "category": current_category or "Other",
                "topic": topic,
                "dri": dri,
                "forum": forum,
                "status": status,
                "decision_doc": decision_doc,
                "decision_makers": decision_makers,
                "decision_outcome": decision_outcome,
                "post": post,
                "order": len(decisions)
            })
    
    return decisions

def parse_software_review(docx_path):
    """
    Parse the Software review document and extract items by section
    Returns list of dicts with: section_type, content, is_new, order
    """
    doc = Document(docx_path)
    items = []
    current_section = None
    order = 0
    
    # Section markers
    wins_markers = ["🏆 Wins", "Wins"]
    exec_summary_markers = ["🚀 Exec Summary", "Exec Summary"]
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            continue
            
        # Check for section headers
        if any(marker in text for marker in wins_markers):
            current_section = "wins"
            order = 0
            continue
        elif any(marker in text for marker in exec_summary_markers):
            current_section = "exec_summary"
            order = 0
            continue
        elif "Product Decisions" in text:
            # Stop processing paragraphs when we hit Product Decisions
            break
        
        # Skip if we haven't found a section yet
        if current_section is None:
            continue
            
        # Skip section headers and empty lines
        if text.startswith("📣") or text.startswith("FYIs"):
            continue
        if text.startswith("🗓️ Upcoming Releases"):
            break
        if text.startswith("Portfolio View"):
            break
        if text.startswith("🚩 Leadership Help Needed"):
            break
            
        # Check if this is a content line (starts with bracket or bullet)
        if text.startswith("[") or text.startswith("•") or text.startswith("-") or text.startswith("**"):
            # Check if any run in this paragraph has blue text
            has_blue = any(is_blue_text(run) for run in para.runs)
            
            # Get numbering level for indentation
            indent_level = 0
            numbering_part = para._element.pPr.numPr if para._element.pPr is not None and hasattr(para._element.pPr, 'numPr') else None
            if numbering_part is not None and numbering_part.ilvl is not None:
                doc_level = numbering_part.ilvl.val
                if doc_level >= 2:
                    indent_level = doc_level - 1
            
            # Extract rich text with bold and links
            rich_content = extract_rich_text(para)
            
            items.append({
                "section_type": current_section,
                "content": rich_content,
                "is_new": 1 if has_blue else 0,
                "indent_level": indent_level,
                "order": order
            })
            order += 1
    
    # Extract Product Decisions table
    decisions = parse_product_decisions_table(doc)
    items.extend(decisions)
    
    return items

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: parse_software_review.py <docx_file>", file=sys.stderr)
        sys.exit(1)
    
    docx_path = sys.argv[1]
    items = parse_software_review(docx_path)
    print(json.dumps(items, indent=2))
