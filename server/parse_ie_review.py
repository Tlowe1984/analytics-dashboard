#!/usr/bin/env python3
"""
Parse Experiences & Interfaces Review document (weekly WXX file)
Extracts Wins, Exec Summary, and Decisions for three sections: I+E, AI, Hearing
"""

import sys
import json
import subprocess
import re
from docx import Document
from datetime import datetime, timedelta
from rich_text_parser_v2 import extract_rich_text_with_links
from extract_decision_tables import extract_ie_decisions, extract_ai_decisions, extract_hearing_decisions
from extract_hotspots import extract_hotspots_with_wearables_tag

def get_current_week():
    """Get current week number (ISO week)"""
    return datetime.now().isocalendar()[1]

def find_latest_review_file():
    """Find the most recent WXX Experiences & Interfaces Review file based on modification time"""
    try:
        # List files with metadata in the I+E reviews folder
        result = subprocess.run(
            [
                "rclone", "lsjson",
                "manus_google_drive:Wearables Everything/Reviews (Comment Only)/Software (I+E, AI, Hearing) Reviews/I+E Previous Reviews & Review Notes/",
                "--config", "/home/ubuntu/.gdrive-rclone.ini"
            ],
            capture_output=True,
            text=True,
            timeout=120
        )
        
        if result.returncode != 0:
            print(f"rclone error: {result.stderr}", file=sys.stderr)
            return None, None
        
        # Parse JSON and find WXX Experiences & Interfaces Review files
        files_data = json.loads(result.stdout)
        review_files = []
        
        for item in files_data:
            name = item.get('Name', '')
            # Flexible match — handles all known naming variants:
            #   "W09 Experiences & Interfaces Review.docx"                   (original)
            #   "W09 Software (I+E, AI, Hearing) Canonical Program Review.docx" (new canonical)
            #   "WK09 Experiences & Interfaces Review.docx"                  (WK prefix)
            #   "W9 Software (I+E, AI, Hearing) Review.docx"                 (single digit)
            if re.match(
                r'W(K)?\d+\s+(Experiences\s*[&+]\s*Interfaces\s+Review'
                r'|Software\s*\(I\+E[^)]*\).*Review)',
                name, re.IGNORECASE
            ):
                mod_time = item.get('ModTime', '')
                review_files.append({
                    'name': name,
                    'modified': mod_time,
                    'id': item.get('ID', '')
                })
        
        if not review_files:
            # Broad fallback: any .docx containing both "Software" and "Review" (or I+E and Review)
            print("Primary pattern matched nothing — trying broad Software/I+E Review fallback...", file=sys.stderr)
            for item in files_data:
                name = item.get('Name', '')
                if re.search(r'(Software|I\+E|Experiences).*Review.*\.docx', name, re.IGNORECASE):
                    mod_time = item.get('ModTime', '')
                    review_files.append({'name': name, 'modified': mod_time, 'id': item.get('ID', '')})
        
        if not review_files:
            print("No Software/I+E review files found in folder", file=sys.stderr)
            return None, None, None
        
        # Sort by modification time (most recent first)
        review_files.sort(key=lambda x: x['modified'], reverse=True)
        latest_file = review_files[0]['name']
        latest_modified = review_files[0]['modified']
        latest_id = review_files[0].get('id', '')
        
        print(f"Found latest review file: {latest_file} (modified {latest_modified})", file=sys.stderr)
        return latest_file, latest_modified, latest_id
    
    except Exception as e:
        print(f"Error finding latest review file: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return None, None, None

def download_review_file(filename):
    """Download the review file from Google Drive"""
    try:
        # Delete old file first to force fresh download
        local_path = f"/tmp/{filename}"
        subprocess.run(["rm", "-f", local_path], check=False)
        
        result = subprocess.run(
            [
                "rclone", "copy",
                f"manus_google_drive:Wearables Everything/Reviews (Comment Only)/Software (I+E, AI, Hearing) Reviews/I+E Previous Reviews & Review Notes/{filename}",
                "/tmp/",
                "--config", "/home/ubuntu/.gdrive-rclone.ini",
                "--ignore-times",
                "--no-check-certificate"
            ],
            capture_output=True,
            text=True,
            timeout=120
        )
        
        if result.returncode != 0:
            print(f"rclone error: {result.stderr}", file=sys.stderr)
            return None, None
        
        # Validate file week (must be current or last week)
        # Get file modification time
        file_stat_result = subprocess.run(
            ['stat', '-c', '%y', f"/tmp/{filename}"],
            capture_output=True,
            text=True
        )
        if file_stat_result.returncode == 0:
            file_mtime = file_stat_result.stdout.strip().split()[0] + 'T' + file_stat_result.stdout.strip().split()[1].split('.')[0] + 'Z'
            validation_result = subprocess.run(
                ['python3', '/home/ubuntu/analytics-dashboard/server/validate_week.py', file_mtime],
                capture_output=True,
                text=True
            )
            try:
                validation = json.loads(validation_result.stdout)
                if not validation['valid']:
                    print(f"⚠️ WARNING: {validation['message']}", file=sys.stderr)
                    print(f"File may contain outdated data. Consider updating the source document.", file=sys.stderr)
                else:
                    print(f"✅ File validation passed: {validation['message']}", file=sys.stderr)
            except:
                pass  # Skip validation if it fails
        
        return f"/tmp/{filename}"
    
    except Exception as e:
        print(f"Error downloading review file: {e}", file=sys.stderr)
        return None

def parse_section(doc, start_idx, end_idx, section_name):
    """Parse a section of the document and extract Wins, Exec Summary, Decisions"""
    
    wins = []
    exec_summary = []
    decisions = []
    help_needed = []
    
    current_subsection = None
    current_item_parts = []  # Collect parts of current win/exec item
    
    for i in range(start_idx, end_idx):
        para = doc.paragraphs[i]
        text = para.text.strip()
        rich_text = extract_rich_text_with_links(para)  # Get markdown-formatted text
        
        if not text:
            continue
        
        # Detect subsections (flexible patterns to handle spacing/capitalization variations)
        # Wins section: 🏆 + (Wins|Launches) + optional (Async|/)
        if re.search(r'🏆\s*(Wins?|Launches?)', text, re.IGNORECASE):
            # Save previous item if any
            if current_item_parts:
                if current_subsection == 'wins':
                    wins.append('\n'.join(current_item_parts))
                elif current_subsection == 'exec_summary':
                    exec_summary.append('\n'.join(current_item_parts))
                current_item_parts = []
            current_subsection = 'wins'
            continue
        # Exec Summary section: 🚀 + Exec Summary + optional (FYIs|Async)
        elif re.search(r'(🚀|📣)\s*(Exec\s+Summary|FYIs?)', text, re.IGNORECASE):
            # Save previous item if any
            if current_item_parts:
                if current_subsection == 'wins':
                    wins.append('\n'.join(current_item_parts))
                elif current_subsection == 'exec_summary':
                    exec_summary.append('\n'.join(current_item_parts))
                current_item_parts = []
            current_subsection = 'exec_summary'
            continue
        # Decisions section
        elif re.search(r'(Product\s+)?Decisions?\s*\[', text, re.IGNORECASE):
            # Save previous item if any
            if current_item_parts:
                if current_subsection == 'wins':
                    wins.append('\n'.join(current_item_parts))
                elif current_subsection == 'exec_summary':
                    exec_summary.append('\n'.join(current_item_parts))
                current_item_parts = []
            current_subsection = 'decisions'
            continue
        # Help Needed / Flags section (flexible pattern to catch all variations)
        # Matches: "🚩 Help Needed", "Help Needed/ Flag for Leadership", "🚩 Leadership Help Needed", "🚩 Flag", etc.
        elif re.search(r'(🚩|Help\s+Needed|Flag).*?(Help\s+Needed|Flag|Leadership)', text, re.IGNORECASE):
            # Save previous item if any
            if current_item_parts:
                if current_subsection == 'wins':
                    wins.append('\n'.join(current_item_parts))
                elif current_subsection == 'exec_summary':
                    exec_summary.append('\n'.join(current_item_parts))
                current_item_parts = []
            current_subsection = 'help_needed'
            continue
        
        # Extract content based on current subsection
        if current_subsection == 'wins':
            # Check if this is a top-level item (category header)
            is_top_level = (
                text.startswith('[') or 
                (len(text) > 20 and ':' in text[:50]) or
                (len(text) < 50 and not para.style.name.startswith('List Bullet 2'))
            )
            
            if is_top_level and current_item_parts:
                # Save previous item and start new one
                wins.append('\n'.join(current_item_parts))
                current_item_parts = [rich_text]
            else:
                # Add to current item (sub-bullet or continuation)
                current_item_parts.append(rich_text)
        
        elif current_subsection == 'exec_summary':
            # Check if this is a top-level item
            is_top_level = (
                text.startswith('[') or 
                (len(text) > 20 and ':' in text[:50]) or
                (len(text) < 50 and not para.style.name.startswith('List Bullet 2'))
            )
            
            if is_top_level and current_item_parts:
                # Save previous item and start new one
                exec_summary.append('\n'.join(current_item_parts))
                current_item_parts = [rich_text]
            else:
                # Add to current item
                current_item_parts.append(rich_text)
        
        elif current_subsection == 'help_needed':
            if text.startswith('[') or text.startswith('-') or text.startswith('•') or text.startswith('○'):
                help_needed.append(rich_text.lstrip('-•○ ').strip())
        
        elif current_subsection == 'decisions':
            # Decisions are often in tables, capture text that looks like decision items
            if text and len(text) > 10:
                decisions.append(rich_text)
    
    # Save last item if any
    if current_item_parts:
        if current_subsection == 'wins':
            wins.append('\n'.join(current_item_parts))
        elif current_subsection == 'exec_summary':
            exec_summary.append('\n'.join(current_item_parts))
    
    # Process items to detect and strip [wearables-tag]
    def process_items_for_wearables_tag(items):
        processed = []
        for item in items:
            # Check for [wearables-tag] (case insensitive)
            has_tag = bool(re.search(r'\[wearables-tag\]', item, re.IGNORECASE))
            # Strip the tag from content
            cleaned_item = re.sub(r'\[wearables-tag\]', '', item, flags=re.IGNORECASE).strip()
            processed.append({
                'content': cleaned_item,
                'is_wearables_tag': 1 if has_tag else 0
            })
        return processed
    
    return {
        'section': section_name,
        'wins': process_items_for_wearables_tag(wins[:20]),  # Limit to top 20 items
        'exec_summary': process_items_for_wearables_tag(exec_summary[:30]),  # Limit to top 30 items
        'help_needed': process_items_for_wearables_tag(help_needed[:10]),  # Limit to top 10 items
        'decisions': process_items_for_wearables_tag(decisions[:20])  # Limit to top 20 items
    }

def main():
    print("Finding latest I+E review file...", file=sys.stderr)
    
    # Find latest review file
    latest_file, latest_modified, latest_file_id = find_latest_review_file()
    if not latest_file:
        print("Failed to find review file", file=sys.stderr)
        print("[]")
        return 1
    
    # Write source metadata for the sync script to pick up
    file_url = f"https://docs.google.com/document/d/{latest_file_id}/edit" if latest_file_id else ''
    with open('/tmp/software_source_meta.json', 'w') as _f:
        json.dump({'filename': latest_file, 'modified': latest_modified or '', 'file_url': file_url}, _f)
    
    # Download the file
    print(f"Downloading {latest_file}...", file=sys.stderr)
    local_path = download_review_file(latest_file)
    if not local_path:
        print("Failed to download review file", file=sys.stderr)
        print("[]")
        return 1
    
    # Parse the document
    print(f"Parsing {local_path}...", file=sys.stderr)
    try:
        doc = Document(local_path)
        
        # Find section boundaries
        sections = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if re.match(r'^(Experiences?\s*&\s*Interfaces?|AI|Hearing)\s*$', text, re.IGNORECASE):
                sections.append({
                    'name': text,
                    'index': i
                })
        
        if len(sections) < 3:
            print(f"[PARSER] Warning: Only found {len(sections)} section(s), expected 3 (Experiences & Interfaces, AI, Hearing)", file=sys.stderr)
            # Log all headings found so format changes are diagnosable
            print("[PARSER] All short headings found in document:", file=sys.stderr)
            for i, para in enumerate(doc.paragraphs):
                t = para.text.strip()
                if t and len(t) < 60 and any(c.isupper() for c in t[:5]):
                    print(f"  [{i}] {t!r}", file=sys.stderr)
        
        # Extract decisions from tables
        print("Extracting decision tables...", file=sys.stderr)
        ie_decisions = extract_ie_decisions(doc)
        ai_decisions = extract_ai_decisions(doc)
        hearing_decisions = extract_hearing_decisions(doc)
        
        # Extract hotspots with wearables-tag
        print("Extracting wearables-tagged hotspots...", file=sys.stderr)
        hotspots = extract_hotspots_with_wearables_tag(doc)
        
        decisions_map = {
            'Experiences & Interfaces': ie_decisions,
            'AI': ai_decisions,
            'Hearing': hearing_decisions
        }
        
        # Parse each section
        results = []
        for idx, section in enumerate(sections):
            start = section['index']
            end = sections[idx + 1]['index'] if idx + 1 < len(sections) else len(doc.paragraphs)
            
            section_data = parse_section(doc, start, end, section['name'])
            
            # Add structured decisions
            section_decisions = decisions_map.get(section['name'], [])
            section_data['structured_decisions'] = section_decisions
            
            # Add hotspots to Experiences & Interfaces section only
            if section['name'] == 'Experiences & Interfaces':
                section_data['hotspots'] = hotspots
            
            results.append(section_data)
            
            print(f"Parsed {section['name']}: {len(section_data['wins'])} wins, {len(section_data['exec_summary'])} exec items, {len(section_decisions)} decisions", file=sys.stderr)
        
        # Output as JSON
        print(json.dumps(results, indent=2))
        return 0
    
    except Exception as e:
        print(f"Error parsing document: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        print("[]")
        return 1

if __name__ == "__main__":
    sys.exit(main())
