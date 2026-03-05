#!/usr/bin/env python3
"""
Parse AI WX Review document
Extracts Wins/Launches [Async], Exec Summary [Async], and This Week's Decisions sections
"""

import sys
import json
from docx import Document
from docx.table import Table
from rich_text_parser_v2 import extract_rich_text

def parse_ai_review(docx_path):
    """
    Parse AI Review document to extract:
    - Wins/Launches: bullet points under 🏆Wins/Launches [Async] or Wins/Launches heading
    - Exec Summary: bullet points under 🚀Exec Summary [Async] or Exec Summary heading
    - Decisions: table rows from This Week's Decisions or AI Decisions table
    Returns structured data with rich text formatting preserved.
    """
    doc = Document(docx_path)
    
    items = []
    current_section = None
    order = 0
    pending_wearables_tag = False  # Propagate tag to sub-bullets when tagged item is a heading
    wearables_heading_indent = 0  # Indent level of the tagged heading, propagate to deeper levels
    import re
    
    # Process all document elements (paragraphs and tables)
    for element in doc.element.body:
        # Handle paragraphs
        if element.tag.endswith('p'):
            # Find the paragraph object
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            
            if para is None:
                continue
                
            text = para.text.strip()
            
            if not text:
                continue
            
            # Check if this is a section header (with or without emoji)
            text_lower = text.lower()
            
            # Look for Wins/Launches section (🏆Wins/Launches [Async] or just Wins/Launches)
            if ('wins' in text_lower or 'launches' in text_lower) and ('[async]' in text_lower or len(text) < 30):
                current_section = 'wins'
                continue
            
            # Look for Exec Summary section (🚀Exec Summary [Async] or just Exec Summary)
            if 'exec summary' in text_lower and ('[async]' in text_lower or len(text) < 40):
                current_section = 'exec_summary'
                continue
            
            # Look for Help Needed section header - skip this section entirely
            if 'help needed' in text_lower or 'flag for leadership' in text_lower:
                current_section = 'help_needed'
                continue
            
            # Look for Decisions section header
            if 'decisions' in text_lower and ('this week' in text_lower or 'ai decisions' in text_lower):
                current_section = 'decisions'
                continue
            
            # Skip if we're in decisions or help_needed section waiting for table
            if current_section in ['decisions', 'help_needed']:
                continue
            
            # Skip if we haven't found a section yet
            if current_section is None:
                # Log unrecognised headings so format changes are visible in sync logs
                if len(text) < 60 and any(c.isupper() for c in text[:5]):
                    print(f"[PARSER] Unrecognised heading skipped: {text!r}", file=sys.stderr)
                continue
            
            # Extract rich text content for Wins and Exec Summary
            rich_text = extract_rich_text(para)
            
            if not rich_text or rich_text.strip() == '':
                continue
            
            # Detect if this is a "new" item (blue text)
            is_new = False
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    r, g, b = run.font.color.rgb
                    # Blue text detection
                    if b > 150 and r < 100 and g < 100:
                        is_new = True
                        break
            
            # Detect and strip [wearables-tag]
            is_wearables_tag = '[wearables-tag]' in rich_text.lower()
            if is_wearables_tag:
                # Remove [wearables-tag] (case insensitive)
                rich_text = re.sub(r'\[wearables-tag\]', '', rich_text, flags=re.IGNORECASE).strip()
            
            # Detect indent level
            indent_level = 0
            if para.paragraph_format.left_indent:
                indent_inches = para.paragraph_format.left_indent.inches
                indent_level = int(indent_inches / 0.5)
            
            # Propagate wearables tag to sub-bullets when a tagged item is just a heading (ends with ':')
            # e.g. "[wearables-tag] Harmony:" should tag all following indented sub-bullets
            # Propagation continues for ALL paragraphs indented deeper than the tagged heading
            if not is_wearables_tag and pending_wearables_tag:
                if indent_level > wearables_heading_indent:
                    is_wearables_tag = True  # Sub-bullet under a tagged heading
                else:
                    pending_wearables_tag = False  # Back to same/higher level, stop propagating
            
            if is_wearables_tag:
                stripped = rich_text.rstrip()
                if stripped.endswith(':') and len(stripped) < 40:
                    pending_wearables_tag = True  # This is a heading, propagate to sub-bullets
                    wearables_heading_indent = indent_level  # Remember this heading's indent level
                # Don't reset pending_wearables_tag for non-heading tagged items — keep propagating siblings
            
            items.append({
                'section_type': current_section,
                'content': rich_text,
                'is_new': is_new,
                'is_wearables_tag': is_wearables_tag,
                'indent_level': indent_level,
                'order': order
            })
            
            order += 1
        
        # Handle tables (for Decisions section)
        elif element.tag.endswith('tbl'):
            # Find the table object
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break
            
            if table is None:
                continue
            
            # Only process tables with "AI Decisions" or "This Week's Decisions" in the header
            if len(table.rows) == 0:
                continue
            
            first_cell_text = table.rows[0].cells[0].text.strip()
            # Check if this is a decisions table
            is_decisions_table = ('AI Decisions' in first_cell_text or 
                                'This Week' in first_cell_text or
                                'Decisions' in first_cell_text)
            
            if not is_decisions_table:
                continue  # Skip tables that aren't the main Decisions table
            
            # Parse table rows (skip title row and header row)
            for i, row in enumerate(table.rows):
                if i == 0 or i == 1:  # Skip title row (0) and header row (1)
                    continue
                
                cells = row.cells
                if len(cells) < 6:  # Need at least 6 columns
                    continue
                
                # Extract cell contents with rich text
                # Columns: DRI, Forum, Status, Decision Doc, Decision Makers, Decision outcome/Steer, Post
                dri = cells[0].text.strip()
                forum = cells[1].text.strip()
                status = cells[2].text.strip()
                decision_doc = extract_rich_text(cells[3].paragraphs[0]) if cells[3].paragraphs else cells[3].text.strip()
                decision_makers = cells[4].text.strip()
                decision_outcome = extract_rich_text(cells[5].paragraphs[0]) if cells[5].paragraphs else cells[5].text.strip()
                post = cells[6].text.strip() if len(cells) > 6 else ''
                
                # Skip empty rows
                if not dri or dri == '':
                    continue
                
                # Skip section header rows (e.g., "FYI Sub-Pillar & Cross-Pillar Decisions")
                if 'sub-pillar' in dri.lower() or 'cross-pillar' in dri.lower():
                    continue
                
                # Detect and strip [wearables-tag] from decision outcome
                is_wearables_tag = '[wearables-tag]' in decision_outcome.lower()
                if is_wearables_tag:
                    import re
                    decision_outcome = re.sub(r'\[wearables-tag\]', '', decision_outcome, flags=re.IGNORECASE).strip()
                
                # Create decision item with table structure
                items.append({
                    'section_type': 'decisions',
                    'content': decision_outcome,  # Main content is the decision outcome
                    'dri': dri,
                    'forum': forum,
                    'status': status,
                    'decision_doc': decision_doc,
                    'decision_makers': decision_makers,
                    'post': post,
                    'is_new': False,
                    'is_wearables_tag': is_wearables_tag,
                    'indent_level': 0,
                    'order': order
                })
                
                order += 1
    
    return {
        'items': items,
        'total_count': len(items)
    }

def main():
    if len(sys.argv) < 2:
        print(json.dumps({'error': 'Usage: parse_ai_review.py <docx_path>'}))
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    try:
        result = parse_ai_review(docx_path)
        print(json.dumps(result, indent=2))
    except Exception as e:
        print(json.dumps({'error': str(e)}), file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main()
