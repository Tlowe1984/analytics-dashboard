#!/usr/bin/python3.11
"""
Parser for Wearables Decision Canonical document
Extracts decisions from Consolidated Summary table and filters for last month
"""
import sys
import json
from docx import Document
from datetime import datetime, timedelta
from rich_text_parser_v2 import extract_rich_text_from_cell

def parse_week_number(week_str):
    """Parse week string like 'W49 2025' or 'WW51' to get week and year"""
    if not week_str:
        return None, None
    
    week_str = week_str.strip().upper()
    
    # Handle formats like "W49 2025" or "WW51 2025"
    if ' ' in week_str:
        parts = week_str.split()
        week_part = parts[0].replace('W', '').replace('WW', '')
        year_part = parts[1] if len(parts) > 1 else str(datetime.now().year)
    else:
        # Handle formats like "W49" or "WW51"
        week_part = week_str.replace('W', '').replace('WW', '')
        year_part = str(datetime.now().year)
    
    try:
        week_num = int(week_part)
        year = int(year_part)
        return week_num, year
    except:
        return None, None

def is_within_last_20_weeks(week_str):
    """Check if the week is within the last 20 weeks"""
    week_num, year = parse_week_number(week_str)
    if week_num is None or year is None:
        return False
    
    # Get current week number
    current_date = datetime.now()
    current_week = current_date.isocalendar()[1]
    current_year = current_date.year
    
    # Calculate week difference
    if year == current_year:
        week_diff = current_week - week_num
    elif year == current_year - 1:
        # Handle year boundary (e.g., W51 2025 when current is W2 2026)
        week_diff = current_week + (52 - week_num)
    else:
        return False
    
    # Include decisions from last 20 weeks
    return 0 <= week_diff <= 20

def parse_decisions(doc_path):
    """Extract decisions from the Consolidated Summary table"""
    doc = Document(doc_path)
    
    if len(doc.tables) == 0:
        print("No tables found in document", file=sys.stderr)
        return []
    
    # First table is the Consolidated Summary
    table = doc.tables[0]
    decisions = []
    
    for i, row in enumerate(table.rows[1:]):  # Skip header row
        # Extract rich text from cells
        cells = [extract_rich_text_from_cell(cell).strip() for cell in row.cells]
        
        # Skip if not enough columns, header row, or no decision outcome
        if len(cells) < 5 or cells[0] == "DRI" or cells[0] == "Decisions" or "**DRI**" in cells[0]:
            continue
        
        # Skip if no decision outcome (last column)
        if not cells[4] or len(cells[4].strip()) < 10:
            continue
        
        week_str = cells[3]
        
        # Filter for last 20 weeks
        if not is_within_last_20_weeks(week_str):
            continue
        
        decision = {
            "dri": cells[0],
            "forum": cells[1],
            "status": cells[2],
            "week": week_str,
            "decision_outcome": cells[4]
        }
        decisions.append(decision)
    
    # Sort by week (most recent first)
    decisions.sort(key=lambda x: parse_week_number(x["week"]) or (0, 0), reverse=True)
    
    return decisions

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: parse_decisions.py <docx_path>", file=sys.stderr)
        sys.exit(1)
    
    doc_path = sys.argv[1]
    decisions = parse_decisions(doc_path)
    
    print(json.dumps(decisions, indent=2))
