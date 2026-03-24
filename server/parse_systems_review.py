#!/usr/bin/env python3
"""
Parse Wearables Systems Review document (weekly WK## file)
Extracts Wins, Exec Summary, and Help Needed sections
"""

import sys
import json
import subprocess
import re
from docx import Document
from datetime import datetime
from rich_text_parser_v2 import extract_rich_text
from extract_systems_wearables import extract_wearables_tagged_systems

def find_latest_systems_review_file():
    """Find the most recent WK## Wearables Systems Review file based on modification time"""
    try:
        # List files with metadata in the Systems Software Reviews/Archive folder
        result = subprocess.run(
            [
                "rclone", "lsjson",
                "manus_google_drive:Wearables Everything/Reviews (Comment Only)/Systems Software Reviews/Archive/",
                "--config", "/home/ubuntu/.gdrive-rclone.ini"
            ],
            capture_output=True,
            text=True,
            timeout=120
        )
        
        if result.returncode != 0:
            print(f"rclone error: {result.stderr}", file=sys.stderr)
            return None, None
        
        # Parse JSON and find WK## Wearables Systems Review files
        files_data = json.loads(result.stdout)
        review_files = []
        
        for item in files_data:
            name = item.get('Name', '')
            # Flexible match — handles all known naming variants:
            #   "Wearables Systems Review-WK09-2026.docx"
            #   "Wearables Systems Review WK09 2026.docx"
            #   "Wearables Systems Review-W09-2026.docx"   (no K)
            #   "Wearables Systems Review W09 2026.docx"   (no K, space)
            #   "Wearables Systems Review WK9 2026.docx"   (single digit)
            if re.match(r'Wearables\s+Systems\s+Review[-\s]+WK?\d+[-\s]+\d{4}\.docx', name, re.IGNORECASE):
                mod_time = item.get('ModTime', '')
                review_files.append({
                    'name': name,
                    'modified': mod_time,
                    'id': item.get('ID', '')
                })
        
        if not review_files:
            # Broad fallback: any .docx containing "Systems" and "Review" in the folder
            print("Primary pattern matched nothing — trying broad Systems Review fallback...", file=sys.stderr)
            for item in files_data:
                name = item.get('Name', '')
                if re.search(r'Systems.*Review.*\.docx', name, re.IGNORECASE):
                    mod_time = item.get('ModTime', '')
                    review_files.append({'name': name, 'modified': mod_time, 'id': item.get('ID', '')})
        
        if not review_files:
            print("No Systems review files found in archive folder", file=sys.stderr)
            return None, None, None
        
        def week_sort_key(f):
            """Sort by (year, week_number) extracted from filename — NOT by modification time.
            This prevents an older week's file from being picked just because it was edited recently.
            e.g. WK12-2026 > WK11-2026 regardless of which file was touched last."""
            name = f['name']
            # Extract year (default 2026 if not present)
            year_match = re.search(r'(\d{4})', name)
            year = int(year_match.group(1)) if year_match else 2026
            # Extract week number — handles WK09, WK9, W09, W9
            wk_match = re.search(r'WK?(\d+)', name, re.IGNORECASE)
            week = int(wk_match.group(1)) if wk_match else 0
            return (year, week)

        # Sort by (year, week number) descending — highest week wins
        review_files.sort(key=week_sort_key, reverse=True)
        latest_file = review_files[0]['name']
        latest_modified = review_files[0]['modified']
        latest_id = review_files[0].get('id', '')
        
        print(f"Found latest Systems review file: {latest_file} (week sort key: {week_sort_key(review_files[0])}, modified {latest_modified})", file=sys.stderr)
        return latest_file, latest_modified, latest_id
    
    except Exception as e:
        print(f"Error finding latest Systems review file: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return None, None, None

def download_systems_review_file(filename):
    """Download the Systems review file from Google Drive"""
    try:
        # Clear cache first
        subprocess.run(
            ["rm", "-f", f"/tmp/{filename}"],
            capture_output=True,
            timeout=5
        )
        
        result = subprocess.run(
            [
                "rclone", "copy",
                f"manus_google_drive:Wearables Everything/Reviews (Comment Only)/Systems Software Reviews/Archive/{filename}",
                "/tmp/",
                "--config", "/home/ubuntu/.gdrive-rclone.ini",
                "--ignore-times",
                "--no-check-certificate"
            ],
            capture_output=True,
            text=True,
            timeout=120
        )
        
        if result.returncode != 0:
            print(f"rclone error: {result.stderr}", file=sys.stderr)
            return None
        
        return f"/tmp/{filename}"
    
    except Exception as e:
        print(f"Error downloading Systems review file: {e}", file=sys.stderr)
        return None

def parse_systems_review(docx_path):
    """
    Parse Wearable Systems Review document to extract Wins, Exec Summary, and Help Needed sections.
    Returns a list of items with section type, content, and is_new flag (for blue text).
    """
    doc = Document(docx_path)
    items = []
    current_section = None
    order = 0
    
    import re as _re
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            continue
            
        # Check for section headers using flexible regex patterns
        # Wins: any heading containing "Wins" or "Launches" (with optional emoji/async)
        if _re.search(r'(🏆|^)\s*(Wins?|Launches?)\s*(\[Async\])?\s*$', text, _re.IGNORECASE):
            current_section = "wins"
            continue
        # Exec Summary: any heading containing "Exec Summary" or "FYIs" (with optional emoji/async)
        elif _re.search(r'(🚀|📣|^)\s*(Exec\s+Summary|FYIs?)\s*(\[Async\])?\s*$', text, _re.IGNORECASE):
            current_section = "exec_summary"
            continue
        # Help Needed: any heading mentioning Help Needed or Flag for Leadership
        elif _re.search(r'(🆘|🚩)?\s*(Help\s+Needed|Flag\s+(for\s+)?Leadership)', text, _re.IGNORECASE):
            current_section = "help_needed"
            continue
        
        # Skip if we haven't entered a section yet
        if current_section is None:
            # Log unrecognised headings so format changes are visible in sync logs
            if len(text) < 60 and any(c.isupper() for c in text[:5]):
                print(f"[PARSER] Unrecognised heading skipped: {text!r}", file=sys.stderr)
            continue
        
        # Check if this is a bullet point or content line (not a sub-header)
        # Skip lines that look like sub-headers (short, all caps, or contain only emojis)
        if len(text) < 10 or text.isupper() or all(c in "⚠️🔴✅🎯🎉🆘🏆🚀 " for c in text):
            continue
        
        # Detect blue text (new information)
        is_new = 0
        for run in para.runs:
            if run.font.color and run.font.color.type == 1:  # RGB color
                rgb = run.font.color.rgb
                # Check if it's blue-ish (R < 100, G < 150, B > 150)
                # RGBColor is a tuple-like (r, g, b)
                if rgb[0] < 100 and rgb[1] < 150 and rgb[2] > 150:
                    is_new = 1
                    break
        
        # Get numbering level for indentation
        # Level 0: Section headers
        # Level 1: Main bullets (should be flush left, indent_level=0)
        # Level 2+: Sub-bullets (should be indented, indent_level=1+)
        indent_level = 0
        numbering_part = para._element.pPr.numPr if para._element.pPr is not None and hasattr(para._element.pPr, 'numPr') else None
        if numbering_part is not None and numbering_part.ilvl is not None:
            doc_level = numbering_part.ilvl.val
            # Map document levels to UI indent levels
            # doc_level 0,1 -> indent_level 0 (flush left)
            # doc_level 2+ -> indent_level 1+ (indented)
            if doc_level >= 2:
                indent_level = doc_level - 1
        
        # Extract rich text with bold and links
        rich_content = extract_rich_text(para)
        
        # Detect and strip [wearables-tag] for inline detection
        is_wearables_tag = 0
        if '[wearables-tag]' in rich_content.lower():
            is_wearables_tag = 1
            # Remove [wearables-tag] (case insensitive)
            import re
            rich_content = re.sub(r'\[wearables-tag\]', '', rich_content, flags=re.IGNORECASE).strip()
        
        items.append({
            "section_type": current_section,
            "content": rich_content,
            "is_new": is_new,
            "is_wearables_tag": is_wearables_tag,
            "indent_level": indent_level,
            "order": order
        })
        order += 1
    
    return items

if __name__ == "__main__":
    # Find latest weekly archive file
    print("Finding latest Systems review file...", file=sys.stderr)
    latest_file, latest_modified, latest_file_id = find_latest_systems_review_file()
    
    if not latest_file:
        print("Error: Could not find latest Systems review file", file=sys.stderr)
        sys.exit(1)
    
    # Download the file
    print(f"Downloading {latest_file}...", file=sys.stderr)
    docx_path = download_systems_review_file(latest_file)
    
    if not docx_path:
        print("Error: Could not download Systems review file", file=sys.stderr)
        sys.exit(1)
    
    # Parse the document
    print(f"Parsing {latest_file}...", file=sys.stderr)
    items = parse_systems_review(docx_path)
    
    # Extract wearables-tagged items
    print("Extracting wearables-tagged items...", file=sys.stderr)
    doc = Document(docx_path)
    wearables_items = extract_wearables_tagged_systems(doc)
    
    # Add wearables-tagged items to the output (already classified as wins or exec_summary)
    for idx, item in enumerate(wearables_items):
        items.append({
            "section_type": item['section_type'],
            "content": item['content'],
            "is_new": 0,
            "is_wearables_tag": 1,  # Mark as wearables-tagged item
            "indent_level": 0,
            "order": len(items) + idx
        })
    
    # Write source metadata for the sync script to pick up
    file_url = f"https://docs.google.com/document/d/{latest_file_id}/edit" if latest_file_id else ''
    with open('/tmp/systems_source_meta.json', 'w') as _f:
        json.dump({'filename': latest_file, 'modified': latest_modified or '', 'file_url': file_url}, _f)
    
    # Output JSON to stdout
    print(json.dumps(items, indent=2))
